from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "outputs"
ASSET_DIR = ROOT / ".codex_tmp" / "entropy_guide_assets"
OUTPUT_PATH = OUTPUT_DIR / "Rand_AI_Entropy_Implementation_Guide.docx"
FIGURE_PATH = ASSET_DIR / "entropy_gap_examples.png"

NAVY = "173D64"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "24364B"
MUTED = "66768A"
PALE_BLUE = "E8EEF5"
LIGHT_BLUE = "F2F7FB"
LIGHT_GRAY = "F2F4F7"
GOLD = "D89A2B"
PALE_GOLD = "FFF6DF"
RED = "A63B32"
GREEN = "2D7A56"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def configure_document(doc: Document) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        left = p.add_run("RAND AI  ·  TECHNICAL GUIDE")
        set_run_font(left, size=9, color=MUTED, bold=True)
        right = p.add_run("\tENTROPY STRATEGY")
        set_run_font(right, size=9, color=MUTED, bold=True)

    for footer in (section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        add_page_number(p)


def create_list_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_list_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, italic: bool = False, after: float | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    apply_list_numbering(p, doc._randai_bullet_num_id)
    run = p.add_run(text)
    set_run_font(run)


def add_numbered(doc: Document, title: str, explanation: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_list_numbering(p, num_id)
    title_run = p.add_run(title)
    set_run_font(title_run, bold=True, color=DARK_BLUE)
    explanation_run = p.add_run(f" {explanation}")
    set_run_font(explanation_run)


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT_BLUE, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=accent)
    text_run = p.add_run(text)
    set_run_font(text_run, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, heading in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(heading)
        set_run_font(run, size=10.5, color=NAVY, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, size=10.25)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf")
    return ImageFont.truetype(str(path), size)


def draw_entropy_figure() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (1400, 570), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((6, 6, 1394, 564), radius=28, outline="#CED9E5", width=3, fill="#F8FAFC")
    draw.text((60, 32), "Two gap patterns, two entropy scores", font=font(34, True), fill="#173D64")
    draw.text(
        (60, 78),
        "Each example contains six numbers. The sixth gap wraps from the last number back to the first.",
        font=font(21),
        fill="#52667D",
    )

    examples = [
        {
            "x": 55,
            "title": "Balanced around the circle",
            "numbers": [1, 9, 17, 25, 33, 41],
            "gaps": [8, 8, 8, 8, 8, 9],
            "score": "99.94%",
            "fill": "#E8F3EE",
            "accent": "#2D7A56",
        },
        {
            "x": 720,
            "title": "Clustered in one area",
            "numbers": [1, 2, 3, 4, 5, 6],
            "gaps": [1, 1, 1, 1, 1, 44],
            "score": "27.56%",
            "fill": "#FBECE9",
            "accent": "#A63B32",
        },
    ]
    for example in examples:
        x = example["x"]
        draw.rounded_rectangle((x, 130, x + 620, 525), radius=22, fill=example["fill"], outline=example["accent"], width=3)
        draw.text((x + 28, 154), example["title"], font=font(25, True), fill=example["accent"])
        draw.text((x + 28, 201), "Numbers", font=font(18, True), fill="#52667D")
        for index, number in enumerate(example["numbers"]):
            cx = x + 62 + index * 94
            cy = 260
            draw.ellipse((cx - 29, cy - 29, cx + 29, cy + 29), fill="#FFFFFF", outline="#173D64", width=3)
            text = str(number)
            bbox = draw.textbbox((0, 0), text, font=font(20, True))
            draw.text((cx - (bbox[2] - bbox[0]) / 2, cy - 13), text, font=font(20, True), fill="#173D64")
        draw.text((x + 28, 313), "Circular gaps", font=font(18, True), fill="#52667D")
        for index, gap in enumerate(example["gaps"]):
            bx = x + 31 + index * 94
            by = 352
            fill = "#FFF6DF" if index == 5 else "#FFFFFF"
            draw.rounded_rectangle((bx, by, bx + 70, by + 50), radius=10, fill=fill, outline="#D89A2B", width=2)
            text = str(gap)
            bbox = draw.textbbox((0, 0), text, font=font(20, True))
            draw.text((bx + 35 - (bbox[2] - bbox[0]) / 2, by + 12), text, font=font(20, True), fill="#7B5717")
        draw.text((x + 28, 442), "Normalized entropy", font=font(18, True), fill="#52667D")
        draw.text((x + 330, 428), example["score"], font=font(38, True), fill=example["accent"])
    image.save(FIGURE_PATH, quality=95)


def add_figure(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(FIGURE_PATH), width=Inches(6.45))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(8)
    text = caption.add_run(
        "Figure 1. A balanced set divides the 1–49 circle into similar pieces; a clustered set leaves one very large wrap-around gap."
    )
    set_run_font(text, size=9.5, color=MUTED, italic=True)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    draw_entropy_figure()
    doc = Document()
    configure_document(doc)
    doc._randai_bullet_num_id = create_list_numbering(doc, "bullet")

    # Cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(88)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("RAND AI IMPLEMENTATION GUIDE")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("How Rand AI Uses Entropy")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("A visual guide to the prediction strategy")
    set_run_font(run, size=15, color=DARK_BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.55)
    lead.paragraph_format.right_indent = Inches(0.55)
    lead.paragraph_format.space_after = Pt(70)
    run = lead.add_run(
        "Entropy asks a simple question: are six lottery numbers spread around 1–49 in a balanced way, or crowded into one part of the range?"
    )
    set_run_font(run, size=15, color=INK, italic=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    run = meta.add_run("Written for high-school readers")
    set_run_font(run, size=11, color=MUTED, bold=True)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Based on the current Rand AI source implementation · July 2026")
    set_run_font(run, size=10, color=MUTED)

    add_page_break(doc)

    # Page 2
    doc.add_heading("1. The big idea", level=1)
    add_body(
        doc,
        "In information science, entropy measures how evenly something is distributed. Rand AI applies that idea to the spaces between six lottery numbers.",
    )
    add_callout(
        doc,
        "Pizza-slice analogy",
        "Imagine the numbers 1 through 49 arranged in a circle. The six chosen numbers cut that circle into six slices. Similar slice sizes mean high entropy. One huge slice and several tiny slices mean low entropy.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_body(doc, "The strategy is interested in draw shape, not in whether a number is “good” or “bad” by itself.")
    add_bullet(doc, "High entropy: the six circular gaps are similar in size.")
    add_bullet(doc, "Low entropy: some gaps are tiny while another gap is very large.")
    add_bullet(doc, "The result is normalized to a percentage from 0% to 100%.")
    add_figure(doc)
    add_callout(
        doc,
        "Important",
        "Entropy does not prove that a combination is more likely to win. It is a descriptive pattern used to rank candidates from historical data.",
        fill="FBECE9",
        accent=RED,
    )

    add_page_break(doc)

    # Page 3
    doc.add_heading("2. How one draw becomes an entropy percentage", level=1)
    add_body(doc, "Rand AI performs the following steps for every completed draw.")
    entropy_steps_num_id = create_list_numbering(doc, "decimal")
    add_numbered(doc, "Sort the six numbers.", "For example: 4, 7, 15, 26, 37, 49.", entropy_steps_num_id)
    add_numbered(doc, "Measure five forward gaps.", "Subtract each number from the next one: 3, 8, 11, 11, and 12.", entropy_steps_num_id)
    add_numbered(
        doc,
        "Add the circular gap.",
        "Wrap from 49 back to 4: (49 + 4) - 49 = 4. The six gaps always add to 49.",
        entropy_steps_num_id,
    )
    add_numbered(
        doc,
        "Convert gaps to shares.",
        "Divide every gap by 49. A gap of 8 becomes 8/49, or about 0.163.",
        entropy_steps_num_id,
    )
    add_numbered(
        doc,
        "Calculate Shannon entropy.",
        "Use each share in the standard entropy formula, then add the six contributions.",
        entropy_steps_num_id,
    )
    add_numbered(
        doc,
        "Normalize the answer.",
        "Divide by log2(6), the maximum possible entropy for six parts, and multiply by 100.",
        entropy_steps_num_id,
    )
    add_callout(
        doc,
        "Formula used in the code",
        "H = -sum(p_i x log2(p_i));    Entropy percent = H / log2(6) x 100",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    doc.add_heading("Why divide by log2(6)?", level=2)
    add_body(
        doc,
        "There are six gaps. Entropy is greatest when all six shares are equal. Dividing by that theoretical maximum changes the result into an easy 0–100 scale.",
    )
    add_table(
        doc,
        ["Example draw", "Six circular gaps", "Entropy", "Meaning"],
        [
            ["1, 9, 17, 25, 33, 41", "8, 8, 8, 8, 8, 9", "99.94%", "Almost perfectly balanced"],
            ["4, 7, 15, 26, 37, 49", "3, 8, 11, 11, 12, 4", "94.14%", "Fairly balanced"],
            ["1, 2, 3, 4, 5, 6", "1, 1, 1, 1, 1, 44", "27.56%", "Strongly clustered"],
        ],
        [2300, 2500, 1250, 3310],
    )
    add_body(
        doc,
        "Technical note: the implementation uses differences between selected numbers. The circular total is 49. This is slightly different from a “space” definition that counts only unselected numbers between two selected values.",
        italic=True,
    )

    add_page_break(doc)

    # Page 4
    doc.add_heading("3. What Rand AI remembers from history", level=1)
    add_body(
        doc,
        "After calculating a draw’s entropy, Rand AI gives that same entropy percentage to every number that appeared in the draw. It updates the history before creating the prediction for the next draw.",
    )
    add_callout(
        doc,
        "Example",
        "If draw 100 contains number 17 and has 94% entropy, number 17 receives 94 points in its entropy history for that appearance.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    doc.add_heading("Two history records are kept for each number", level=2)
    add_table(
        doc,
        ["Record", "What is stored", "Why it matters"],
        [
            [
                "Entropy total",
                "The sum of entropy percentages from every draw containing that number.",
                "Dividing by appearances gives the number’s average draw entropy.",
            ],
            [
                "High-entropy hits",
                "How many appearances occurred in draws with entropy at least 92%.",
                "Dividing by appearances gives the high-entropy share.",
            ],
        ],
        [1900, 3600, 3860],
    )
    add_body(
        doc,
        "The code also tracks appearances and the current gap: how many draws have passed since a number last appeared.",
    )
    doc.add_heading("The three-part candidate score", level=2)
    add_table(
        doc,
        ["Part", "Weight", "Question answered"],
        [
            ["Average entropy", "55%", "When this number appeared, how balanced were its draws on average?"],
            ["High-entropy share", "30%", "How often did it appear in draws scoring at least 92%?"],
            ["Current gap", "15%", "How overdue is it now? Full gap credit is reached at 28 draws."],
        ],
        [2500, 1200, 5660],
    )
    add_callout(
        doc,
        "Raw score",
        "0.55 x average entropy + 0.30 x high-entropy share in percent + 0.15 x overdue score",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    add_page_break(doc)

    # Page 5
    doc.add_heading("4 | Worked candidate example", level=1)
    add_body(
        doc,
        "Suppose number 17 has appeared 100 times in the history used by the prediction engine.",
    )
    add_table(
        doc,
        ["Input", "Value", "Contribution"],
        [
            ["Average entropy", "88%", "88 x 0.55 = 48.4"],
            ["High-entropy share", "25%", "25 x 0.30 = 7.5"],
            ["Current gap", "14 draws", "(14 / 28 x 100) x 0.15 = 7.5"],
            ["Raw total", "—", "48.4 + 7.5 + 7.5 = 63.4"],
        ],
        [2700, 1800, 4860],
    )
    add_body(
        doc,
        "The overdue part is capped. A current gap of 28 or more receives 100% of the available gap component, which is 15 points.",
    )
    add_callout(
        doc,
        "If a number has never appeared",
        "Rand AI starts its average entropy at 50%, its high-entropy share at 0%, and still allows the current-gap component to contribute.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    doc.add_heading("From raw scores to a Top-6 prediction", level=2)
    ranking_steps_num_id = create_list_numbering(doc, "decimal")
    add_numbered(
        doc,
        "Score all 49 numbers.",
        "Every candidate receives the same three-part calculation.",
        ranking_steps_num_id,
    )
    add_numbered(
        doc,
        "Rescale to 0–1.",
        "The lowest raw score becomes 0 and the highest becomes 1. Values between them are scaled proportionally.",
        ranking_steps_num_id,
    )
    add_numbered(
        doc,
        "Sort from highest to lowest.",
        "If scores tie, the number with the larger current gap goes first. If that also ties, the lower number goes first.",
        ranking_steps_num_id,
    )
    add_numbered(
        doc,
        "Take the first six.",
        "Those numbers are displayed as the Entropy strategy’s Top-6 prediction.",
        ranking_steps_num_id,
    )
    add_callout(
        doc,
        "Edge case",
        "If every raw score is identical, the rescaling function assigns 0 to all candidates; gap and number order then break the ranking ties.",
        fill=LIGHT_GRAY,
        accent=MUTED,
    )

    add_page_break(doc)

    # Page 6
    doc.add_heading("5 | Where users see the result", level=1)
    add_body(
        doc,
        "The strategy is registered as Entropy and described in the application as “Structural gap-entropy history with overdue adjustment.”",
    )
    add_bullet(doc, "Predictions ranks all 49 numbers and marks the six highest candidates.")
    add_bullet(doc, "Each candidate can show its average entropy and high-entropy share.")
    add_bullet(doc, "Strategy Effectiveness compares completed Top-6 results with a random benchmark.")
    add_bullet(doc, "Entropy can also act as one expert inside Rand AI’s combined strategies.")
    doc.add_heading("Three entropy readings that should not be confused", level=2)
    add_table(
        doc,
        ["Location", "What it measures"],
        [
            [
                "Entropy prediction strategy",
                "Historical circular-gap shape for draws containing each candidate, plus a current-gap adjustment. This guide focuses on this implementation.",
            ],
            [
                "Possible Draw entropy status",
                "The circular-gap entropy of the six numbers currently selected in a plan. It labels the set as Clustered, Balanced, or Correlated.",
            ],
            [
                "Randomness diagnostics",
                "How evenly all 49 numbers have appeared in the entire dataset. This uses frequency entropy, not the six-gap candidate score.",
            ],
        ],
        [2500, 6860],
    )
    doc.add_heading("What the strategy can and cannot say", level=2)
    add_bullet(doc, "It can summarize whether past draws containing a number tended to have balanced spacing.")
    add_bullet(doc, "It can combine that pattern with a small overdue-number adjustment.")
    add_bullet(doc, "It cannot change the lottery’s rules or make independent random events predictable.")
    add_bullet(doc, "A high score is a ranking signal, not a winning probability.")
    add_bullet(doc, "The 92% threshold, 28-draw cap, and 55/30/15 weights are design choices and should be tested rather than treated as natural laws.")
    add_callout(
        doc,
        "Responsible interpretation",
        "Use entropy as one transparent feature among several. Judge it with walk-forward results, compare it with random selections, and avoid claiming certainty from a pattern score.",
        fill="FBECE9",
        accent=RED,
    )

    add_page_break(doc)

    # Page 7
    doc.add_heading("6. Quick reference", level=1)
    add_table(
        doc,
        ["Item", "Implementation value"],
        [
            ["Lottery range", "1 through 49"],
            ["Numbers per draw", "6"],
            ["Gap type", "Six forward differences, including one circular wrap-around gap"],
            ["Entropy formula", "Shannon entropy, base 2"],
            ["Normalization", "Divide by log2(6), then multiply by 100"],
            ["High-entropy threshold", "92%"],
            ["Score weights", "55% average entropy; 30% high-entropy share; 15% current gap"],
            ["Gap cap", "Full overdue credit at 28 draws"],
            ["Final score scale", "Min-max scaled across all 49 candidates to 0–1"],
            ["Tie-breakers", "Larger current gap, then lower number"],
            ["Output", "A ranking of 49 numbers and a Top-6 selection"],
        ],
        [2850, 6510],
    )
    doc.add_heading("Plain-language glossary", level=2)
    add_body(doc, "Candidate: one of the numbers 1–49 being ranked.", bold_lead="Candidate:")
    add_body(doc, "Circular gap: the forward distance between neighboring selected numbers, including the wrap from the largest number back to the smallest.", bold_lead="Circular gap:")
    add_body(doc, "Entropy: a score for how evenly the six gaps share the full circle.", bold_lead="Entropy:")
    add_body(doc, "High-entropy draw: a draw whose normalized entropy is at least 92%.", bold_lead="High-entropy draw:")
    add_body(doc, "Current gap: the number of draws since a candidate last appeared.", bold_lead="Current gap:")
    add_body(doc, "Min-max scaling: changing a set of scores so the smallest is 0 and the largest is 1.", bold_lead="Min-max scaling:")
    add_page_break(doc)
    doc.add_heading("7. Implementation map", level=1)
    add_body(
        doc,
        "The main logic is in src/rand_ai/strategy_prediction.py:",
    )
    add_bullet(doc, "_gap_entropy_percent — calculates normalized circular-gap entropy.")
    add_bullet(doc, "_StrategyState.remember — records entropy history for drawn numbers.")
    add_bullet(doc, "_StrategyState._entropy_scores — builds the 55/30/15 raw score.")
    add_bullet(doc, "_scale_scores and _ranking_from_scores — normalize and order candidates.")
    add_bullet(doc, "build_prediction_suites — updates history and produces each next-draw prediction.")
    add_callout(
        doc,
        "One-sentence summary",
        "Rand AI’s Entropy strategy favors numbers that have often appeared in evenly spaced draws, especially when those numbers are currently overdue.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.core_properties.title = "How Rand AI Uses Entropy"
    doc.core_properties.subject = "High-school-level implementation guide for the Rand AI Entropy prediction strategy"
    doc.core_properties.author = "Rand AI"
    doc.core_properties.keywords = "Rand AI, entropy, Shannon entropy, lottery, prediction strategy, implementation"
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
