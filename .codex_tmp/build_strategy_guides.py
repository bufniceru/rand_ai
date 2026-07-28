from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"C:\code_py\rand_ai")
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SKILL_SCRIPTS = Path(
    r"C:\Users\Floyd\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.723.12215\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


# compact_reference_guide preset
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
INK = "20252B"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "7A5A00"
PALE_GOLD = "FFF7DD"
GREEN = "246B4A"
PALE_GREEN = "EAF5EF"
RED = "9B1C1C"
PALE_RED = "FBEDED"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=9.3, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_border(paragraph, *, color: str, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), color)
        p_bdr.append(edge)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, color: str = "CAD3DC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = tc_borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            tc_borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_repeat_row_safe(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_real_numbering(document: Document, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


class Guide:
    def __init__(self, title: str, short_title: str, subtitle: str) -> None:
        self.doc = Document()
        self.title = title
        self.short_title = short_title
        self.subtitle = subtitle
        self.bullet_num = add_real_numbering(self.doc, ordered=False)
        self.decimal_num = add_real_numbering(self.doc, ordered=True)
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Twips(PAGE_WIDTH_DXA)
        section.page_height = Twips(PAGE_HEIGHT_DXA)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.font.size = Pt(11)
        normal.font.color.rgb = rgb(INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        normal.paragraph_format.widow_control = True

        heading_specs = {
            "Heading 1": (16, BLUE, 18, 10),
            "Heading 2": (13, BLUE, 14, 7),
            "Heading 3": (12, DARK_BLUE, 10, 5),
        }
        for name, (size, color, before, after) in heading_specs.items():
            style = styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = rgb(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style.paragraph_format.keep_together = True

        self.doc.core_properties.title = self.title
        self.doc.core_properties.subject = self.subtitle
        self.doc.core_properties.author = "Rand AI documentation"
        self.doc.core_properties.keywords = (
            "Rand AI, lottery analysis, prediction strategy, educational guide"
        )

        self._set_running_header_footer(section)

    def _set_running_header_footer(self, section) -> None:
        header = section.header
        p = header.paragraphs[0]
        p.text = ""
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
        )
        left = p.add_run(self.short_title)
        set_run_font(left, size=9, color=MUTED, bold=True)
        right = p.add_run("\tRAND AI · LEARNING GUIDE")
        set_run_font(right, size=8.5, color=MUTED)

        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        label = p.add_run("Page ")
        set_run_font(label, size=9, color=MUTED)
        add_page_field(p)

    def cover(self, kicker: str, audience: str) -> None:
        for _ in range(5):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(kicker.upper())
        set_run_font(r, size=10.5, color=GOLD, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(self.title)
        set_run_font(r, size=30, color=NAVY, bold=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(28)
        r = p.add_run(self.subtitle)
        set_run_font(r, size=15, color=DARK_BLUE)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(90)
        r = p.add_run(
            "A plain-language guide to the mathematics, program logic, "
            "and correct reading of the results"
        )
        set_run_font(r, size=10.5, color=MUTED, italic=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Prepared for")
        set_run_font(r, size=9, color=MUTED, bold=True)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(audience)
        set_run_font(r, size=11.5, color=NAVY, bold=True)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Implementation reference: Rand AI · 25 July 2026")
        set_run_font(r, size=9.5, color=MUTED)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def chapter(
        self,
        title: str,
        intro: str | None = None,
        *,
        new_page: bool = True,
    ) -> None:
        heading = self.doc.add_heading(title, level=1)
        heading.paragraph_format.page_break_before = new_page
        if intro:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(intro)
            set_run_font(r, size=11.2, color=DARK_BLUE, italic=True)

    def heading(self, text: str, level: int = 2) -> None:
        self.doc.add_heading(text, level=level)

    def body(self, text: str) -> None:
        self.doc.add_paragraph(text)

    def body_runs(self, parts: list[tuple[str, bool, bool]]) -> None:
        p = self.doc.add_paragraph()
        for text, bold, italic in parts:
            r = p.add_run(text)
            set_run_font(r, bold=bold, italic=italic)

    def bullet(self, text: str) -> None:
        p = self.doc.add_paragraph()
        apply_num(p, self.bullet_num)
        r = p.add_run(text)
        set_run_font(r)

    def numbered(self, text: str, *, restart: bool = False) -> None:
        if restart:
            self.decimal_num = add_real_numbering(self.doc, ordered=True)
        p = self.doc.add_paragraph()
        apply_num(p, self.decimal_num)
        r = p.add_run(text)
        set_run_font(r)

    def callout(
        self,
        label: str,
        text: str,
        *,
        tone: str = "blue",
    ) -> None:
        palette = {
            "blue": (CALLOUT, BLUE),
            "gold": (PALE_GOLD, GOLD),
            "green": (PALE_GREEN, GREEN),
            "red": (PALE_RED, RED),
        }
        fill, accent = palette[tone]
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(9)
        shade_paragraph(p, fill)
        set_paragraph_border(p, color=accent, size=6)
        r = p.add_run(f"{label}: ")
        set_run_font(r, color=accent, bold=True)
        r = p.add_run(text)
        set_run_font(r, color=INK)

    def equation(self, formula: str, explanation: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(4)
        shade_paragraph(p, LIGHT_GRAY)
        r = p.add_run(formula)
        set_run_font(r, name="Consolas", size=10.5, color=NAVY, bold=True)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(9)
        r = p.add_run(explanation)
        set_run_font(r, size=9.5, color=MUTED, italic=True)

    def table(
        self,
        headers: list[str],
        rows: list[list[str]],
        widths_dxa: list[int],
        *,
        font_size: float = 9.3,
    ) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        set_repeat_table_header(table.rows[0])
        set_repeat_row_safe(table.rows[0])
        for idx, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY)
            set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
            set_cell_border(table.rows[0].cells[idx])
        for row_idx, values in enumerate(rows):
            row = table.add_row()
            set_repeat_row_safe(row)
            for col_idx, value in enumerate(values):
                set_cell_text(row.cells[col_idx], value)
                for run in row.cells[col_idx].paragraphs[0].runs:
                    run.font.size = Pt(font_size)
                if row_idx % 2 == 1:
                    set_cell_shading(row.cells[col_idx], "FAFBFC")
                set_cell_border(row.cells[col_idx])
        apply_table_geometry(
            table,
            widths_dxa,
            table_width_dxa=CONTENT_WIDTH_DXA,
            indent_dxa=TABLE_INDENT_DXA,
            cell_margins_dxa=CELL_MARGINS_DXA,
        )
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    def contents(self, items: list[tuple[str, str]]) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run("CONTENTS AND LEARNING MAP")
        set_run_font(r, size=20, color=NAVY, bold=True)
        self.body(
            "You can read this guide from start to finish, or jump directly to the "
            "question you need to answer. Each chapter separates three ideas: what "
            "the strategy measures, how the program computes it, and what the "
            "displayed result does—and does not—mean."
        )
        self.table(
            ["Chapter", "What you will learn"],
            [[name, purpose] for name, purpose in items],
            [2700, 6660],
        )
        self.callout(
            "The most important reading rule",
            "A high score means “higher according to this model and this history.” "
            "It does not mean “certain to be drawn.” Lottery draws remain random.",
            tone="gold",
        )

    def save(self, path: Path) -> None:
        # Ask Word to refresh fields such as page numbers when the file opens.
        settings = self.doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
        self.doc.save(path)


def build_freshness() -> Path:
    path = OUTPUT_DIR / "Freshness_Strategy_Guide.docx"
    g = Guide(
        "Freshness Strategy",
        "Freshness Strategy",
        "How Rand AI learns from the number of draws since a number last appeared",
    )
    g.cover(
        "Rand AI strategy guide",
        "Readers with high-school mathematics; no programming background required",
    )
    g.contents(
        [
            ("1. The idea in everyday language", "Understand gap, hit, exposure, and support."),
            ("2. The computation", "Follow the training table, smoothing formula, and worked examples."),
            ("3. The implementation", "See the program’s walk-forward sequence and ranking rules."),
            ("4. Reading the results", "Interpret Gap, Hit probability, Score, Rank, and Top 6."),
            ("5. Limits and responsible use", "Know what the model can describe and what it cannot predict."),
            ("Reference", "Use the checklist, glossary, and compact algorithm summary."),
        ]
    )

    g.chapter(
        "1. The idea in everyday language",
        "Freshness asks a simple historical question: after a number has waited this "
        "many draws, how often did a number in the same waiting state appear next?",
    )
    g.heading("1.1 One number, two possible outcomes")
    g.body(
        "For every candidate number from 1 to 49, each draw has only two outcomes. "
        "The number is drawn, which we can write as D, or it is not drawn, written "
        "as !D. A standard draw contains 6 numbers, so every draw creates 49 "
        "number-opportunities: 6 are hits and 43 are misses."
    )
    g.table(
        ["Term", "Plain-language meaning"],
        [
            ["Gap", "How many complete draws have passed since the number last appeared."],
            ["Exposure", "One historical chance for a number in a particular gap state."],
            ["Hit", "An exposure whose next outcome was D."],
            ["Support", "The number of exposures used for that exact gap."],
            ["Estimated probability", "The smoothed historical hit rate for that gap."],
        ],
        [2200, 7160],
    )
    g.heading("1.2 What a gap means")
    g.body(
        "Gap 0 means the number appeared in the most recent known draw. Gap 1 means "
        "it missed the most recent draw but appeared one draw earlier. Gap 2 means "
        "two complete draws have passed since its last appearance, and so on."
    )
    g.table(
        ["Recent history for one number", "Current gap", "Reason"],
        [
            ["…  !D  !D  D", "0", "The latest draw is D."],
            ["…  !D  D  !D", "1", "One draw has passed since D."],
            ["…  D  !D  !D", "2", "Two draws have passed since D."],
            ["Never observed yet", "Number of known draws", "All known draws have been misses."],
        ],
        [3200, 1400, 4760],
    )
    g.callout(
        "Fresh does not mean “must be overdue”",
        "The model does not assume that a long absence creates a debt that the "
        "lottery must repay. It only measures what happened historically after each "
        "exact gap.",
        tone="gold",
    )

    g.chapter(
        "2. The computation",
        "The strategy builds one shared lookup table: exact gap → hits and exposures.",
    )
    g.heading("2.1 Learning one draw without looking into the future")
    g.numbered(
        "Before reading the outcome of a historical draw, calculate the current gap "
        "for every number from 1 to 49.",
        restart=True,
    )
    g.numbered(
        "Check the draw. Add 1 exposure to the row for that gap. If the number was "
        "drawn, also add 1 hit."
    )
    g.numbered(
        "After learning the outcome, update the last-seen position of every drawn "
        "number. The state is now ready for the next draw."
    )
    g.body(
        "This order matters. If the program updated last-seen before it learned the "
        "outcome, the answer would leak into the input. Rand AI deliberately learns "
        "first and observes second."
    )
    g.heading("2.2 The shared gap table")
    g.body(
        "Freshness pools the evidence from all 49 numbers. Number 7 at gap 3 and "
        "number 31 at gap 3 contribute to the same gap-3 row. This gives more "
        "evidence than building 49 separate gap tables, but it also means the model "
        "treats the meaning of gap 3 as common to every number."
    )
    g.table(
        ["Exact gap", "Historical hits", "Exposures (support)", "Raw hit rate"],
        [
            ["0", "530", "4,300", "530 ÷ 4,300 = 12.33%"],
            ["1", "465", "3,900", "465 ÷ 3,900 = 11.92%"],
            ["3", "18", "120", "18 ÷ 120 = 15.00%"],
            ["14", "1", "2", "1 ÷ 2 = 50.00% (very uncertain)"],
        ],
        [1400, 1900, 2300, 3760],
    )
    g.body(
        "The numbers in this table are teaching examples. The real table is built "
        "from the currently loaded draw history."
    )
    g.heading("2.3 Why smoothing is necessary")
    g.body(
        "A rare gap can produce a dramatic raw percentage from only one or two "
        "examples. Rand AI therefore adds a small prior worth two imaginary "
        "exposures at the ordinary lottery rate."
    )
    g.equation(
        "smoothed rate = (hits + 2 × 6/49) ÷ (exposures + 2)",
        "6/49 ≈ 12.24% is the neutral chance for one specified number in a 6-of-49 draw.",
    )
    g.table(
        ["Example", "Raw rate", "Smoothed result", "What changed"],
        [
            [
                "Gap 3: 18 hits / 120 exposures",
                "15.00%",
                "(18 + 0.2449) ÷ 122 = 14.95%",
                "Large support, so the adjustment is small.",
            ],
            [
                "Gap 14: 1 hit / 2 exposures",
                "50.00%",
                "(1 + 0.2449) ÷ 4 = 31.12%",
                "Tiny support, so the estimate moves toward 12.24%.",
            ],
            [
                "Unseen gap: 0 / 0",
                "Undefined",
                "0.2449 ÷ 2 = 12.24%",
                "No evidence means the neutral rate.",
            ],
        ],
        [2350, 1200, 3050, 2760],
        font_size=8.8,
    )
    g.callout(
        "Support changes confidence, not just the percentage",
        "Two gap states can show similar probabilities, but the one based on 500 "
        "exposures is more stable than the one based on 8. Always read support "
        "together with probability when support is available.",
        tone="blue",
    )

    g.chapter(
        "3. How Rand AI implements Freshness",
        "The implementation is incremental: it carries only the counts and the last "
        "known appearance of each number as it moves through history.",
    )
    g.heading("3.1 Main state kept in memory")
    g.table(
        ["Program state", "Purpose"],
        [
            [
                "last_seen[1…49]",
                "Stores the draw index where each number was most recently observed.",
            ],
            [
                "freshness_counts[gap] = [hits, exposures]",
                "Stores the shared evidence for every exact gap that has actually occurred.",
            ],
            [
                "BASE_HIT_RATE = 6/49",
                "Supplies the neutral prior used by smoothing.",
            ],
            [
                "PRIOR_STRENGTH = 2",
                "Controls how strongly an estimate with little evidence is pulled toward 6/49.",
            ],
        ],
        [3200, 6160],
    )
    g.heading("3.2 Walk-forward sequence")
    g.table(
        ["Moment", "Known to the model", "Action"],
        [
            ["Before historical draw t", "Draws 1 through t−1", "Compute gaps and learn whether draw t was D or !D."],
            ["After observing draw t", "Draws 1 through t", "Update last-seen positions."],
            ["Prediction for draw t+1", "Still only draws 1 through t", "Look up each current exact gap and rank all 49 numbers."],
            ["Later, when t+1 is known", "Draws 1 through t+1", "Use it as a training outcome for the next step."],
        ],
        [1900, 2600, 4860],
    )
    g.callout(
        "No future leakage",
        "The actual numbers of draw t+1 may be stored later for evaluation, but they "
        "are not used to create the prediction for t+1.",
        tone="green",
    )
    g.heading("3.3 From probability to the displayed score")
    g.body(
        "The Fresh strategy first obtains the smoothed hit probability for each "
        "number’s current gap. It then rescales the 49 probabilities so that the "
        "smallest becomes 0 and the largest becomes 1."
    )
    g.equation(
        "display score = (probability − smallest probability) ÷ (largest − smallest)",
        "If every probability is identical, Rand AI assigns score 0 to all numbers.",
    )
    g.body(
        "This score is a relative position inside the current set of 49 numbers. A "
        "score of 0.80 is not an 80% chance of being drawn. The literal model "
        "estimate is the field named Hit probability."
    )
    g.heading("3.4 Ranking and tie-breaking")
    g.numbered("Higher normalized Fresh score ranks first.", restart=True)
    g.numbered("If scores tie, the larger current gap ranks first.")
    g.numbered("If score and gap both tie, the smaller number ranks first.")
    g.body(
        "The first six ranked numbers become the strategy’s Top 6. These rules make "
        "the result stable and repeatable when several numbers have the same evidence."
    )
    g.heading("3.5 Compact pseudocode")
    g.equation(
        "for each draw: learn(gap → outcome); remember(drawn numbers); "
        "for each number: probability = smoothed_rate(counts[current_gap]); rank",
        "This is a readable summary, not a line-for-line copy of the Python source.",
    )
    g.body(
        "The core implementation is in src/rand_ai/prediction.py. The named Fresh "
        "strategy and score normalization are assembled in "
        "src/rand_ai/strategy_prediction.py."
    )

    g.chapter(
        "4. How to interpret the results",
        "Read the fields in a fixed order: state, evidence-based estimate, relative "
        "score, then rank.",
    )
    g.table(
        ["Displayed item", "Correct interpretation", "Common mistake"],
        [
            ["Gap 7", "Seven complete draws have passed since this number last appeared.", "“It must appear now.”"],
            ["Hit probability 13.80%", "The smoothed historical rate for exact gap 7.", "“The lottery guarantees 13.80%.”"],
            ["Score 0.76", "Relative position between the current minimum and maximum Fresh probabilities.", "“There is a 76% chance.”"],
            ["Rank 4", "Fourth according to Fresh after tie-break rules.", "“Fourth most likely in an objective sense.”"],
            ["Top 6", "The six highest-ranked candidates under this one strategy.", "“A predicted winning ticket.”"],
        ],
        [1800, 4260, 3300],
        font_size=8.7,
    )
    g.heading("4.1 A worked reading example")
    g.table(
        ["Number", "Gap", "Hit probability", "Score", "Interpretation"],
        [
            ["8", "0", "11.90%", "0.10", "Recent appearance; its gap-0 history is near the low end today."],
            ["21", "3", "14.95%", "1.00", "Its current gap maps to the largest Fresh probability among the 49."],
            ["37", "9", "13.10%", "0.45", "Middle of the current Fresh range."],
        ],
        [1000, 900, 1700, 1000, 4760],
        font_size=8.7,
    )
    g.body(
        "In this example, number 21 ranks above number 37 because the shared gap-3 "
        "row has a higher smoothed historical rate than the shared gap-9 row. The "
        "number’s identity is not part of the Fresh probability; its current gap is."
    )
    g.heading("4.2 What a high Fresh score really says")
    g.bullet("The number currently belongs to a gap state that performed relatively well in the loaded history.")
    g.bullet("Its gap state may have strong or weak support; the score alone does not show that distinction.")
    g.bullet("Another strategy can rank the same number very differently because it studies different information.")
    g.bullet("The result is conditional on the chosen history file and its ordering.")
    g.callout(
        "Best short sentence",
        "“Fresh ranks numbers by the historical next-draw rate of their current exact waiting time.”",
        tone="green",
    )

    g.chapter(
        "5. Strengths, limits, and responsible use",
        "Freshness is useful as a transparent historical lens. Transparency does not "
        "turn random draws into a predictable process.",
    )
    g.heading("5.1 Strengths")
    g.bullet("Simple state: the current gap is easy to understand and verify.")
    g.bullet("Walk-forward learning avoids using a future answer in an earlier prediction.")
    g.bullet("Pooling all 49 numbers provides much more evidence for common gap values.")
    g.bullet("Smoothing prevents most tiny samples from becoming immediate 0% or 100% claims.")
    g.bullet("Deterministic ranking makes repeated runs on the same history reproducible.")
    g.heading("5.2 Limits")
    g.bullet("The shared table cannot learn that number 12 behaves differently from number 38 at the same gap.")
    g.bullet("Very large exact gaps remain sparse even after pooling.")
    g.bullet("Min–max scaling can make small probability differences look visually large.")
    g.bullet("Historical association is not a physical cause. Balls do not remember their previous gaps.")
    g.bullet("A fair lottery’s expected chance for one specified number remains 6/49 on each draw.")
    g.heading("5.3 Sensible evaluation")
    g.body(
        "To evaluate Fresh, create each historical prediction using only the draws "
        "that were available at that time. Then count how many of the next draw’s "
        "six numbers were in the predicted Top 6. Compare the long-run average with "
        "a baseline and report the number of tested draws. A short winning streak "
        "is not enough evidence."
    )
    g.callout(
        "Random baseline",
        "If six numbers are selected at random from 49, the expected overlap with "
        "the next six drawn numbers is 6 × 6/49 ≈ 0.735 numbers per draw. Real "
        "results can move above or below this by chance.",
        tone="gold",
    )

    g.chapter(
        "Reference: checklist and glossary",
        "Use this page when explaining a Fresh result to someone else.",
    )
    g.heading("A six-question interpretation checklist")
    for index, item in enumerate([
        "What is the number’s current exact gap?",
        "What smoothed Hit probability belongs to that gap?",
        "How much historical support does that gap have?",
        "Is the displayed value a probability or the normalized 0–1 score?",
        "Was the prediction produced before the target draw was known?",
        "How did the strategy perform over many walk-forward tests, not just one draw?",
    ]):
        g.numbered(item, restart=index == 0)
    g.heading("Glossary")
    g.table(
        ["Word", "Meaning in Freshness"],
        [
            ["Baseline", "The neutral single-number rate 6/49, used as the prior."],
            ["D / !D", "Drawn / not drawn for one number in one draw."],
            ["Exact state", "A precise gap value, not a bucket such as “5–10.”"],
            ["Gap", "Complete draws since the number’s last appearance."],
            ["Hit", "The number was D on the outcome draw."],
            ["Prior", "A small neutral starting belief used before much evidence exists."],
            ["Smoothing", "Combining observed counts with the prior to reduce extremes."],
            ["Support", "The exposure count behind one exact gap estimate."],
            ["Walk-forward", "Train on the past, predict the next step, then move forward."],
        ],
        [1900, 7460],
    )
    g.callout(
        "Final takeaway",
        "Freshness is a gap-conditioned historical frequency model. Its most useful "
        "output is not a promise; it is a clear, reproducible ranking based on how "
        "exact waiting states behaved in the available past.",
        tone="blue",
    )
    g.save(path)
    return path


def build_mkfr() -> Path:
    path = OUTPUT_DIR / "Markov_Freshness_MKFR_Strategy_Guide.docx"
    g = Guide(
        "Markov Freshness (MKFR)",
        "Markov Freshness (MKFR)",
        "How Rand AI learns D / !D transition patterns separately for every number",
    )
    g.cover(
        "Rand AI strategy guide",
        "Readers with high-school mathematics; no programming background required",
    )
    g.contents(
        [
            ("1. The idea in everyday language", "Turn each number’s history into D and !D patterns."),
            ("2. Transition contexts up to order 20", "Understand DD, D!D, !DD, and longer histories."),
            ("3. The computation", "Follow per-number baselines, support, backoff, smoothing, and lift."),
            ("4. The implementation", "See the training sequence, stored tables, and ranking rules."),
            ("5. Reading the results", "Interpret Context probability, Baseline, Lift, Order, and Support."),
            ("6. Limits and responsible use", "Avoid false certainty, pattern myths, and over-reading."),
            ("Reference", "Use the checklist, glossary, and worked calculation."),
        ]
    )

    g.chapter(
        "1. The idea in everyday language",
        "MKFR asks a per-number question: when this same number previously had the "
        "same recent D / !D pattern, how often was its next outcome D?",
    )
    g.heading("1.1 Forty-nine separate binary histories")
    g.body(
        "MKFR builds a separate history for every number from 1 to 49. For number "
        "23, each draw becomes D if 23 appeared and !D if it did not. Number 24 has "
        "its own sequence. The model does not mix their transition counts."
    )
    g.table(
        ["Draw", "Did number 23 appear?", "Symbol stored"],
        [
            ["101", "No", "!D"],
            ["102", "No", "!D"],
            ["103", "Yes", "D"],
            ["104", "No", "!D"],
            ["105", "Yes", "D"],
        ],
        [1800, 4100, 3460],
    )
    g.body(
        "After draw 105, the recent pattern for number 23 is !D !D D !D D. MKFR "
        "can study the last one symbol, the last two symbols, and every longer "
        "suffix up to the last 20 draws."
    )
    g.heading("1.2 What “Markov” means here")
    g.body(
        "A Markov-style model estimates the next outcome from a limited recent "
        "state instead of rereading the entire history as one enormous pattern. "
        "MKFR uses a variable-order state: it begins with short patterns and allows "
        "longer patterns to refine the estimate when enough matching examples exist."
    )
    g.callout(
        "Important",
        "The model treats the recent pattern as useful historical context. It does "
        "not claim that the pattern physically causes the next lottery outcome.",
        tone="gold",
    )

    g.chapter(
        "2. Transition contexts from order 1 to order 20",
        "The order is the number of most recent draws used to describe the current "
        "state of one number.",
    )
    g.heading("2.1 Order 1")
    g.body(
        "Order 1 uses only the latest symbol. There are two possible contexts: D "
        "and !D. The model can therefore estimate P(next D | D) and "
        "P(next D | !D) for each number."
    )
    g.heading("2.2 Order 2")
    g.body(
        "Order 2 uses the last two symbols. There are four possible contexts: DD, "
        "D!D, !DD, and !D!D. For example, D!D means the number was drawn two draws "
        "ago and missed the most recent draw."
    )
    g.table(
        ["Order-2 context", "Older draw", "Most recent draw", "Question asked"],
        [
            ["DD", "D", "D", "After two consecutive appearances, how often was next D?"],
            ["D!D", "D", "!D", "After an appearance then a miss, how often was next D?"],
            ["!DD", "!D", "D", "After a miss then an appearance, how often was next D?"],
            ["!D!D", "!D", "!D", "After two misses, how often was next D?"],
        ],
        [1600, 1500, 1900, 4360],
        font_size=8.8,
    )
    g.heading("2.3 Longer orders and the growth of possible patterns")
    g.body(
        "Every added draw doubles the number of theoretically possible patterns. "
        "Order k has 2ᵏ possible contexts. This growth is why a literal order-20 "
        "table is sparse: many exact 20-draw patterns may never repeat."
    )
    g.table(
        ["Order k", "Draws in context", "Possible D / !D patterns (2ᵏ)"],
        [
            ["1", "1", "2"],
            ["2", "2", "4"],
            ["3", "3", "8"],
            ["5", "5", "32"],
            ["10", "10", "1,024"],
            ["20", "20", "1,048,576"],
        ],
        [1800, 2500, 5060],
    )
    g.callout(
        "What Rand AI stores",
        "It stores only contexts that actually occurred. It does not allocate all "
        "1,048,576 order-20 patterns for all 49 numbers.",
        tone="green",
    )

    g.chapter(
        "3. The computation",
        "MKFR combines four ideas: a personal baseline, exact transition counts, "
        "minimum support, and hierarchical smoothing.",
        new_page=False,
    )
    g.heading("3.1 Each number’s own baseline")
    g.body(
        "Some numbers may have appeared slightly more often than others in the "
        "loaded history purely by chance. MKFR first estimates each number’s normal "
        "historical appearance rate. It smooths that rate toward 6/49 with a prior "
        "strength of 8."
    )
    g.equation(
        "baseline(n) = (appearances(n) + 8 × 6/49) ÷ (draw_count + 8)",
        "The baseline is personal to number n, but starts from the same neutral lottery rate.",
    )
    g.table(
        ["Number", "Appearances / 1,000 draws", "Raw frequency", "Smoothed baseline"],
        [
            ["A", "130", "13.00%", "(130 + 0.9796) ÷ 1,008 = 12.99%"],
            ["B", "115", "11.50%", "(115 + 0.9796) ÷ 1,008 = 11.51%"],
        ],
        [1300, 2800, 1700, 3560],
    )
    g.body(
        "The table is hypothetical. It shows why comparing a context probability "
        "with that same number’s baseline is fairer than ranking only by lifetime "
        "frequency."
    )
    g.heading("3.2 Transition counts")
    g.body(
        "For every number, order, and observed context, the program stores two "
        "counts: how often the next outcome was !D and how often it was D. Their "
        "sum is the context support."
    )
    g.table(
        ["Number", "Order", "Context", "Next !D", "Next D", "Support"],
        [
            ["23", "1", "!D", "710", "96", "806"],
            ["23", "2", "D!D", "82", "13", "95"],
            ["23", "3", "!DD!D", "8", "2", "10"],
            ["23", "4", "D!DD!D", "2", "0", "2 → too little"],
        ],
        [1200, 1000, 1500, 1600, 1400, 2660],
        font_size=8.8,
    )
    g.heading("3.3 Minimum support and backoff")
    g.body(
        "A context must have at least 8 previous next-outcomes before it can change "
        "the probability. If an order has support below 8, MKFR skips it and keeps "
        "the estimate from the latest supported shorter order."
    )
    g.callout(
        "Backoff in one sentence",
        "Use the longest available pattern that has enough evidence, while keeping "
        "the information learned from supported shorter patterns.",
        tone="blue",
    )
    g.heading("3.4 Hierarchical smoothing")
    g.body(
        "Supported contexts are processed from order 1 toward order 20. A longer "
        "context does not completely replace the shorter estimate. Instead, its "
        "observed hit count is combined with the previous probability as a prior "
        "worth 8 opportunities."
    )
    g.equation(
        "new probability = (context hits + 8 × previous probability) ÷ (context support + 8)",
        "The “previous probability” is the baseline first, then the latest supported shorter-order estimate.",
    )
    g.heading("3.5 Worked hierarchical example")
    g.table(
        ["Step", "Evidence", "Calculation", "Result"],
        [
            ["Baseline", "Number’s normal rate", "Given", "12.00%"],
            ["Order 1", "18 D out of 120", "(18 + 8×0.1200) ÷ 128", "14.81%"],
            ["Order 2", "4 D out of 20", "(4 + 8×0.1481) ÷ 28", "18.52%"],
            ["Order 3", "1 D out of 6", "Support 6 < 8, so skip", "18.52%"],
        ],
        [1400, 2400, 3560, 2000],
        font_size=8.8,
    )
    g.body(
        "The selected order is 2 because it is the longest context in this example "
        "with support of at least 8. Order 3 exists but is too weak to update the estimate."
    )

    g.chapter(
        "4. Transition lift and ranking",
        "Lift answers the fairness question: is the current pattern more favorable "
        "than this number’s own normal historical rate?",
    )
    g.heading("4.1 The lift formula")
    g.equation(
        "transition lift = context probability − number’s baseline probability",
        "Rand AI displays the difference in percentage points (pp), not as a percent increase.",
    )
    g.table(
        ["Context probability", "Baseline", "Lift", "Meaning"],
        [
            ["15.20%", "12.00%", "+3.20 pp", "Current pattern raises the estimate above normal."],
            ["11.90%", "12.00%", "−0.10 pp", "Current pattern is almost neutral, slightly below normal."],
            ["8.50%", "12.00%", "−3.50 pp", "Current pattern lowers the estimate below normal."],
        ],
        [2100, 1500, 1600, 4160],
    )
    g.heading("4.2 Why lift reduces number fixation")
    g.body(
        "Suppose number 5 has a high lifetime frequency. Ranking absolute context "
        "probability can keep number 5 near the top even when its current pattern "
        "adds nothing special. Lift subtracts number 5’s own baseline. To rank "
        "high, its current context must improve on what is already normal for number 5."
    )
    g.table(
        ["Number", "Context probability", "Own baseline", "Lift", "Rank signal"],
        [
            ["5", "14.00%", "13.80%", "+0.20 pp", "Small improvement"],
            ["31", "13.20%", "11.40%", "+1.80 pp", "Larger improvement"],
        ],
        [1200, 2200, 1800, 1500, 2660],
    )
    g.body(
        "Although number 5 has the higher absolute context probability, number 31 "
        "has the higher lift and therefore the stronger MKFR ranking signal."
    )
    g.heading("4.3 From lift to the displayed 0–1 score")
    g.body(
        "Rand AI min–max scales the 49 lift values. The lowest current lift becomes "
        "0 and the highest becomes 1. If all lifts are equal, all scores become 0."
    )
    g.equation(
        "display score = (lift − smallest lift) ÷ (largest lift − smallest lift)",
        "A score of 0.90 means high relative lift in this ranking; it does not mean 90% probability.",
    )
    g.heading("4.4 Ranking and tie-breaking")
    g.numbered("Higher normalized transition-lift score ranks first.", restart=True)
    g.numbered("If scores tie, the larger current gap ranks first.")
    g.numbered("If score and gap also tie, the smaller number ranks first.")
    g.body(
        "The first six ranked numbers become MKFR’s Top 6. Tie-breaking makes the "
        "output deterministic; it does not add new statistical evidence."
    )

    g.chapter(
        "5. How Rand AI implements MKFR",
        "The program updates 49 histories and their observed context tables one draw "
        "at a time, always training before it remembers the current outcome.",
    )
    g.heading("5.1 State kept in memory")
    g.table(
        ["Program state", "Purpose"],
        [
            ["mkfr_histories[1…49]", "A rolling D / !D history, maximum length 20, for each number."],
            ["mkfr_transitions[number][order][context]", "Two counts: next !D and next D for one exact context."],
            ["appearances[1…49]", "Lifetime D count used in each number’s baseline."],
            ["draw_count", "Number of draws already remembered."],
            ["MAX_ORDER = 20", "Largest suffix the model may inspect."],
            ["MIN_CONTEXT_SUPPORT = 8", "Minimum evidence required before a context can update probability."],
            ["PRIOR_STRENGTH = 8", "Weight given to the previous probability during smoothing."],
        ],
        [3300, 6060],
        font_size=8.8,
    )
    g.heading("5.2 Training one new draw")
    g.numbered(
        "For each number, label the new draw as target 1 (D) or target 0 (!D).",
        restart=True,
    )
    g.numbered(
        "Before appending that target, read the existing history. Build its suffix "
        "of order 1, then order 2, continuing up to order 20."
    )
    g.numbered(
        "For every available order, increment the stored next-!D or next-D count "
        "for that number and exact context."
    )
    g.numbered(
        "After all transition tables have learned the target, append D or !D to "
        "each number’s rolling history and update appearance counts."
    )
    g.numbered(
        "Now compute the prediction for the following draw from the newly known history."
    )
    g.callout(
        "No future leakage",
        "The current outcome is the target during training, never part of its own "
        "input context. Only after training is it appended for the next prediction.",
        tone="green",
    )
    g.heading("5.3 Compact pseudocode")
    g.equation(
        "for each draw and number: target = D or !D; "
        "for order 1…20: counts[number, order, recent_context][target] += 1; "
        "append target; predict next by baseline → supported contexts → lift → rank",
        "Only observed contexts are stored, using an integer bit pattern internally.",
    )
    g.body(
        "The MKFR constants, training loop, probability backoff, transition lift, "
        "and display details are implemented in src/rand_ai/strategy_prediction.py."
    )
    g.heading("5.4 Why an integer context is used internally")
    g.body(
        "The program stores !D as bit 0 and D as bit 1. A short sequence can then be "
        "used as a compact dictionary key instead of a long text string. This is an "
        "efficiency choice; the interface converts the selected suffix back into a "
        "human-readable D / !D pattern."
    )

    g.chapter(
        "6. How to interpret MKFR results",
        "Read all five detail fields together. No single field is enough to judge a "
        "candidate responsibly.",
    )
    g.table(
        ["Displayed item", "Correct interpretation", "What it is not"],
        [
            ["Context probability 15.20%", "Smoothed estimate after all supported suffixes were processed.", "A guaranteed chance."],
            ["Baseline probability 12.00%", "This number’s smoothed normal historical D rate.", "The same fixed value for every number."],
            ["Transition lift +3.20 pp", "The current context estimate is 3.20 percentage points above baseline.", "A 3.20% relative increase."],
            ["Order 4/20: D!D!DD", "Order 4 is the longest supported suffix used; 20 is the maximum.", "Proof that all 20 draws were useful."],
            ["Context support 11", "The selected order-4 exact context had 11 recorded next-outcomes.", "Total history length."],
            ["Score 0.87", "This lift is near the high end of today’s 49 lifts.", "An 87% chance."],
        ],
        [2150, 4310, 2900],
        font_size=8.4,
    )
    g.heading("6.1 Reading an example from left to right")
    g.callout(
        "Example",
        "Number 23 shows Context probability 15.20%, Baseline 12.00%, Lift +3.20 pp, "
        "Order 4/20: D!D!DD, Support 11, Score 0.87.",
        tone="blue",
    )
    g.numbered(
        "The current four-draw suffix for number 23 is D, !D, !D, D from older to newer.",
        restart=True,
    )
    g.numbered(
        "That exact order-4 suffix has at least 8 examples; here it has 11."
    )
    g.numbered(
        "After hierarchical smoothing, the estimated next-D probability is 15.20%."
    )
    g.numbered(
        "Number 23’s own baseline is 12.00%, so the pattern adds +3.20 percentage points."
    )
    g.numbered(
        "Compared with all 49 current lift values, +3.20 pp scales to 0.87."
    )
    g.numbered(
        "The result supports a relatively high MKFR rank. It still does not make D certain."
    )
    g.heading("6.2 Special cases")
    g.table(
        ["Situation", "What Rand AI does", "How to read it"],
        [
            ["No context reaches support 8", "Uses the number’s baseline; selected order 0; support 0.", "No reliable pattern-specific lift was found."],
            ["Long order is unsupported", "Skips it and keeps the latest supported shorter estimate.", "A shorter pattern supplied the usable evidence."],
            ["Positive probability but negative lift", "Ranks by the negative lift.", "The context is less favorable than this number’s own normal rate."],
            ["Several equal lifts", "Uses gap, then number, as tie-breakers.", "Later tie-breaks are deterministic, not probabilistic."],
        ],
        [2350, 3600, 3410],
        font_size=8.6,
    )

    g.chapter(
        "7. Strengths, limits, and responsible use",
        "MKFR is richer than a simple gap model, but the extra detail creates a "
        "serious risk of sparse data and accidental pattern storytelling.",
        new_page=False,
    )
    g.heading("7.1 Strengths")
    g.bullet("Maintains a separate transition model for each number.")
    g.bullet("Represents recent history directly as D / !D combinations.")
    g.bullet("Uses variable order instead of forcing every number to use all 20 draws.")
    g.bullet("Requires minimum context support before a pattern can affect the estimate.")
    g.bullet("Smooths longer contexts with shorter evidence rather than trusting tiny samples alone.")
    g.bullet("Ranks transition lift above each number’s own baseline, reducing fixation on historically frequent numbers.")
    g.heading("7.2 Limits")
    g.bullet("The number of possible patterns grows exponentially; long contexts are often rare.")
    g.bullet("A support threshold of 8 reduces instability but does not make eight examples strong proof.")
    g.bullet("Repeated testing of many numbers and patterns can reveal impressive-looking coincidences by chance.")
    g.bullet("Min–max scaling can visually magnify small differences in lift.")
    g.bullet("A fair lottery has no physical memory; historical transitions may not persist.")
    g.bullet("The model does not estimate causal mechanisms and cannot guarantee profitable predictions.")
    g.heading("7.3 Sensible evaluation")
    g.body(
        "Evaluate MKFR with walk-forward testing: at every historical point, train "
        "only on earlier draws, create the next Top 6, and compare with the actual "
        "next draw. Report the number of tests, average hits, variability, and a "
        "random baseline. Also check whether particular numbers remain selected for "
        "unusually long streaks; lift should reduce, not completely eliminate, that risk."
    )
    g.callout(
        "Random baseline",
        "Six random selections against six drawn numbers have expected overlap "
        "6 × 6/49 ≈ 0.735 per draw. A small excess in one dataset can be ordinary "
        "sampling noise and should not be described as proof of predictive power.",
        tone="gold",
    )

    g.chapter(
        "Reference: checklist, glossary, and full example",
        "Use this final chapter as a compact explanation sheet.",
        new_page=False,
    )
    g.heading("A seven-question interpretation checklist")
    for index, item in enumerate([
        "Which number is being evaluated, and what is its current D / !D suffix?",
        "What selected order was actually supported?",
        "How many examples support that selected exact context?",
        "What is the smoothed context probability?",
        "What is the same number’s baseline probability?",
        "What is the lift in percentage points, and is it positive or negative?",
        "Is the 0–1 value being read correctly as a relative score rather than a probability?",
    ]):
        g.numbered(item, restart=index == 0)
    g.heading("Glossary")
    g.table(
        ["Word", "Meaning in MKFR"],
        [
            ["Backoff", "Keeping a shorter supported estimate when a longer context lacks evidence."],
            ["Baseline", "One number’s smoothed normal historical appearance probability."],
            ["Context", "The recent D / !D suffix used as the current state."],
            ["D / !D", "Drawn / not drawn for one particular number."],
            ["Lift", "Context probability minus that same number’s baseline, in percentage points."],
            ["Order", "The count of recent draws in a context."],
            ["Prior strength", "How much weight the previous estimate receives during smoothing; MKFR uses 8."],
            ["Support", "The number of recorded next-outcomes following an exact context."],
            ["Transition", "Movement from a context to the next outcome D or !D."],
            ["Variable order", "Using longer contexts only when they have enough evidence."],
        ],
        [1900, 7460],
        font_size=8.8,
    )
    g.heading("Full compact calculation")
    g.table(
        ["Stage", "Input", "Output"],
        [
            ["Personal baseline", "128 D in 1,000 draws; prior strength 8", "(128 + 8×6/49) ÷ 1,008 = 12.80%"],
            ["Order 1", "15 D in 100 matching cases", "(15 + 8×12.80%) ÷ 108 = 14.84%"],
            ["Order 2", "3 D in 16 matching cases", "(3 + 8×14.84%) ÷ 24 = 17.45%"],
            ["Order 3", "1 D in 5 matching cases", "Skipped because support is below 8"],
            ["Lift", "17.45% context − 12.80% baseline", "+4.65 percentage points"],
            ["Ranking", "Compare +4.65 pp with the other 48 lifts", "Normalize to 0–1, then sort"],
        ],
        [1900, 3780, 3680],
        font_size=8.6,
    )
    g.callout(
        "Final takeaway",
        "MKFR is a per-number variable-order transition model. It learns what tended "
        "to follow recent D / !D patterns, protects itself with support and "
        "smoothing, and ranks each pattern’s lift above that number’s own baseline. "
        "It is a disciplined historical comparison—not a guarantee about a random draw.",
        tone="blue",
    )
    g.save(path)
    return path


if __name__ == "__main__":
    for generated in (build_freshness(), build_mkfr()):
        print(generated)
