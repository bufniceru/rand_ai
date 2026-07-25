"""Build the detailed Rand AI prediction calculation guide."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "outputs" / "Rand_AI_Prediction_Calculation_Guide.docx"

# compact_reference_guide preset tokens
FONT = "Calibri"
BODY_SIZE = 11
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "5B6573"
GOLD = "B07A18"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
LINE = "C9D4E0"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float = BODY_SIZE,
    color: str = "172033",
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError("Table columns must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(tbl_header)


def set_paragraph_bottom_border(paragraph, color: str = LINE, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    field_props = OxmlElement("w:rPr")
    field_fonts = OxmlElement("w:rFonts")
    field_fonts.set(qn("w:ascii"), FONT)
    field_fonts.set(qn("w:hAnsi"), FONT)
    field_props.append(field_fonts)
    field_color = OxmlElement("w:color")
    field_color.set(qn("w:val"), MUTED)
    field_props.append(field_color)
    field_size = OxmlElement("w:sz")
    field_size.set(qn("w:val"), "18")
    field_props.append(field_size)
    field_run.append(field_props)
    text = OxmlElement("w:t")
    text.text = "1"
    field_run.append(text)
    fld.append(field_run)
    paragraph._p.append(fld)


def add_numbering(document: Document) -> tuple[int, int, int]:
    """Create fresh Word-native bullet and decimal list instances."""
    numbering = document.part.numbering_part.element

    def fresh_instance(style_name: str) -> int:
        style = document.styles[style_name]
        base_num_id = int(style._element.pPr.numPr.numId.val)
        base_num = numbering.num_having_numId(base_num_id)
        abstract_id = int(base_num.abstractNumId.val)
        abstract = next(
            item
            for item in numbering.findall(qn("w:abstractNum"))
            if int(item.get(qn("w:abstractNumId"))) == abstract_id
        )
        level = abstract.find(qn("w:lvl"))
        p_pr = level.find(qn("w:pPr"))
        if p_pr is None:
            p_pr = OxmlElement("w:pPr")
            level.append(p_pr)
        tabs = p_pr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            p_pr.append(tabs)
        tab = tabs.find(qn("w:tab"))
        if tab is None:
            tab = OxmlElement("w:tab")
            tabs.append(tab)
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "270")
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")

        instance = numbering.add_num(abstract_id)
        instance.add_lvlOverride(ilvl=0).add_startOverride(1)
        return int(instance.numId)

    return (
        fresh_instance("List Bullet"),
        fresh_instance("List Number"),
        fresh_instance("List Number"),
    )


def apply_list_numbering(paragraph, number_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(number_id))
    num_pr.extend((ilvl, num_id))
    p_pr.append(num_pr)


def add_list_item(document: Document, text: str, number_id: int) -> None:
    paragraph = document.add_paragraph()
    apply_list_numbering(paragraph, number_id)
    run = paragraph.add_run(text)
    set_run_font(run)


def add_body(document: Document, text: str, *, bold_lead: str | None = None):
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_equation(document: Document, equation: str, explanation: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.right_indent = Inches(0.22)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shading)
    run = paragraph.add_run(equation)
    set_run_font(run, name="Consolas", size=10.5, color=NAVY, bold=True)
    if explanation:
        run = paragraph.add_run(f"\n{explanation}")
        set_run_font(run, size=9.5, color=MUTED)


def add_callout(document: Document, label: str, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EEF4FA")
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    border.append(left)
    p_pr.append(border)
    run = paragraph.add_run(f"{label}  ")
    set_run_font(run, color=DARK_BLUE, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, color="26384B")


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TABLE_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            paragraph = row.cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(value)
            set_run_font(run, size=9.5, color="26384B")
    configure_table(table, widths)
    # These reference tables are deliberately short enough to stay intact.
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb("172033")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        1: (16, BLUE, 18, 10),
        2: (13, BLUE, 14, 7),
        3: (12, DARK_BLUE, 10, 5),
    }
    for level, (size, color, before, after) in tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(paragraph)


def add_cover(document: Document) -> None:
    # editorial_cover header pattern with a compact technical-guide whitespace override.
    for _ in range(4):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("TECHNICAL REFERENCE GUIDE")
    set_run_font(run, size=10, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Rand AI Combined Prediction")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Exact freshness and left/right proximity calculations")
    set_run_font(run, size=15, color=DARK_BLUE)

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.6)
    summary.paragraph_format.right_indent = Inches(0.6)
    summary.paragraph_format.space_after = Pt(34)
    run = summary.add_run(
        "A detailed explanation of state construction, walk-forward learning, "
        "Bayesian-style smoothing, score combination, ranking, storage, and display."
    )
    set_run_font(run, size=11.5, color=MUTED, italic=True)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(54)
    run = metadata.add_run("Implementation described: Rand AI 0.1.0  |  21 July 2026")
    set_run_font(run, size=10, color=MUTED, bold=True)
    document.add_page_break()


def build_document() -> Document:
    document = Document()
    configure_styles(document)
    configure_page(document)
    bullet_id, decimal_id, decimal_restart_id = add_numbering(document)
    add_cover(document)

    add_heading(document, "1. Purpose and interpretation", 1)
    add_body(
        document,
        "The combined prediction ranks all 49 candidate numbers for the draw immediately following a reference draw. "
        "It combines two empirical signals: freshness, represented by the candidate's exact current absence gap, and "
        "proximity, represented by the exact left/right spaces recorded the last time that candidate appeared."
    )
    add_callout(
        document,
        "Important interpretation",
        "The combined score is a ranking signal derived from historical conditional hit rates. It is not a guarantee, "
        "does not make lottery draws deterministic, and is not the probability that an entire six-number ticket wins.",
    )
    add_heading(document, "Calculation at a glance", 2)
    for text in (
        "At each historical target draw, compute the pre-draw exact gap for every number from 1 through 49.",
        "For numbers seen previously, reuse the exact left/right spacing pair from their most recent appearance.",
        "Record whether each exact state was exposed and whether the candidate was actually drawn.",
        "Convert the accumulated counts into lightly smoothed freshness and proximity hit rates.",
        "Average the two component rates and sort all 49 candidates using deterministic tie-break rules.",
    ):
        add_list_item(document, text, decimal_restart_id)

    add_heading(document, "2. Inputs and exact state representation", 1)
    add_body(
        document,
        "The calculation consumes an ordered Draws collection. Every Draw contains six unique integers in the range "
        "1 through 49. Internally, draws use zero-based positions; the user interface displays draw numbers starting at 1."
    )
    add_table(
        document,
        ["Term", "Implemented meaning", "Value/source"],
        [
            ["Candidate number n", "One number evaluated independently at the current state", "1 to 49"],
            ["Reference draw r", "The draw after which a prediction is produced", "Display numbering starts at 1"],
            ["Target draw", "The draw being predicted", "r + 1"],
            ["Freshness state g", "Exact number of intervening draws since n last appeared", "Non-negative integer"],
            ["Proximity state (L, R)", "Exact circular spaces beside n at its most recent appearance", "Pair of non-negative integers"],
            ["Hit y", "Whether n appears in the historical target draw", "1 if drawn; otherwise 0"],
        ],
        [1900, 4860, 2600],
    )

    add_heading(document, "2.1 Freshness gap", 2)
    add_body(
        document,
        "Before a target draw at internal index t, the engine looks up the internal index s of the candidate's most "
        "recent appearance. The gap excludes both the last appearance and the target draw itself."
    )
    add_equation(
        document,
        "g(n, t) = t                         if n has never appeared\n"
        "g(n, t) = t - s - 1                 otherwise",
        "Example: if a number was last drawn at index 10 and the target is index 15, its exact gap is 15 - 10 - 1 = 4.",
    )
    add_body(
        document,
        "When building a prediction after displayed reference draw r, the code uses reference_index = r - 1. "
        "A number drawn in the reference draw therefore has gap 0 for the next target. An unseen number receives gap r."
    )

    add_heading(document, "2.2 Left and right proximity spaces", 2)
    add_body(
        document,
        "Within each draw, the six values are sorted as a1 < a2 < ... < a6. A space counts the unused lottery "
        "values between neighboring drawn numbers. The number line is circular: the outside space connects a6 back to a1."
    )
    add_equation(
        document,
        "left(a1) = (a1 - 1) + (49 - a6)\n"
        "left(ai) = ai - a(i-1) - 1          for i = 2..6\n"
        "right(ai) = left(a(i+1))            with circular wraparound",
        "For every draw, the six space values sum to 43 because 49 total values minus 6 drawn values leaves 43 unused values.",
    )
    add_body(
        document,
        "For prediction, a candidate retains the L and R values from its most recent appearance. These values are "
        "not guessed for the future draw. They describe the candidate's last observed neighborhood and remain unchanged "
        "until that candidate is drawn again."
    )
    add_callout(
        document,
        "No buckets",
        "Gap 4 is learned separately from gap 5. Likewise, proximity state (2, 5) is learned separately from (2, 6) "
        "and from (5, 2). The implementation never converts these values into ranges or categories.",
    )

    add_heading(document, "3. Walk-forward learning", 1)
    add_body(
        document,
        "The engine learns by replaying the dataset in chronological order. For each eligible historical target draw, "
        "all 49 candidates contribute one freshness exposure. A candidate contributes a proximity exposure only after it "
        "has appeared at least once and therefore has a known exact L/R pair."
    )
    add_table(
        document,
        ["Accumulator", "Key", "Stored counts", "Update rule before target draw"],
        [
            ["freshness_counts", "Exact gap g", "[hits, exposures]", "Add one exposure for every candidate; add one hit if candidate is drawn"],
            ["proximity_counts", "Exact pair (L, R)", "[hits, exposures]", "If L and R are known, add one exposure; add one hit if candidate is drawn"],
        ],
        [1900, 1660, 1900, 3900],
    )
    add_heading(document, "3.1 Per-draw update order", 2)
    for text in (
        "Learn the current draw's outcomes from the states that existed immediately before it.",
        "Observe the current draw: update last_seen, last_left_space, and last_right_space for its six numbers.",
        "Build the prediction for the following draw from the now-current state and all outcomes learned so far.",
    ):
        add_list_item(document, text, decimal_id)
    add_body(
        document,
        "The first stored draw establishes last-seen and proximity state but is not used as a learning target. Starting "
        "with the second draw, each draw is learned before the engine advances to the next prediction."
    )

    add_heading(document, "3.2 Leakage prevention", 2)
    add_body(
        document,
        "For the prediction after reference draw r, the scoring accumulators contain outcomes no later than r. The "
        "actual numbers from target draw r + 1 are copied into CombinedPrediction.actual_numbers only so the historical "
        "grid can show which ranked entries were subsequently drawn. They are never passed into the score calculation."
    )
    add_callout(
        document,
        "Why order matters",
        "Learning target r + 1 before calculating the prediction after r would leak the answer into the model. The "
        "implemented learn -> observe -> predict sequence prevents that error.",
    )

    add_heading(document, "4. From counts to component scores", 1)
    add_body(
        document,
        "Exact states can have limited history. Rand AI therefore applies the same light prior to every component rate. "
        "This stabilizes states with few observations without merging states or changing their raw values."
    )
    add_heading(document, "4.1 Baseline and prior strength", 2)
    add_equation(
        document,
        "p0 = 6 / 49 = 0.1224489796...\nalpha = 2.0",
        "p0 is the unconditional per-number hit rate in a six-of-49 draw. Alpha acts like two baseline exposures.",
    )
    add_heading(document, "4.2 Smoothed exact-state rate", 2)
    add_equation(
        document,
        "smoothed_rate(H, E) = (H + alpha * p0) / (E + alpha)",
        "H is the number of hits observed for the exact state; E is its number of exposures. With E = 0, the result is exactly p0.",
    )
    add_body(
        document,
        "The freshness component looks up the candidate's exact current gap. The proximity component looks up the "
        "candidate's exact retained (L, R) pair. If the pair is unknown or has never been exposed, the proximity component "
        "falls back to the baseline through the same smoothing formula."
    )
    add_equation(
        document,
        "freshness_score(n) = smoothed_rate(freshness_counts[g(n)])\n"
        "proximity_score(n) = smoothed_rate(proximity_counts[(L(n), R(n))])",
    )

    add_heading(document, "5. Combined score and deterministic rank", 1)
    add_body(
        document,
        "Freshness and proximity receive equal weight. No additional features, number-specific multipliers, buckets, "
        "random noise, or manual categories are applied."
    )
    add_equation(
        document,
        "combined_score(n) = (freshness_score(n) + proximity_score(n)) / 2",
    )
    add_heading(document, "5.1 Sorting and tie-breaks", 2)
    add_body(document, "All 49 NumberPrediction records are sorted using the following keys, in order:")
    add_table(
        document,
        ["Priority", "Sort key", "Direction and purpose"],
        [
            ["1", "combined_score", "Descending; strongest overall score first"],
            ["2", "freshness_score", "Descending; resolves equal combined scores"],
            ["3", "proximity_score", "Descending; resolves remaining ties"],
            ["4", "exact gap", "Descending; longer current absence first"],
            ["5", "number", "Ascending; final stable and reproducible tie-break"],
        ],
        [1100, 2500, 5760],
    )
    add_body(
        document,
        "Ranks 1 through 49 are assigned after sorting. top_numbers stores the first six ranked values. The current "
        "Electron grid displays all 49 candidates in this rank order; it does not reorder them numerically from 1 to 49."
    )

    add_heading(document, "6. Worked numerical example", 1)
    worked_intro = add_body(
        document,
        "Assume candidate 17 currently has exact gap g = 4 and retained proximity (L, R) = (2, 5). Suppose the "
        "walk-forward history has the following exact-state counts."
    )
    worked_intro.paragraph_format.keep_with_next = True
    add_table(
        document,
        ["Component", "Exact state", "Hits H", "Exposures E", "Smoothed result"],
        [
            ["Freshness", "g = 4", "8", "50", "(8 + 2 × 6/49) / 52 = 0.158556"],
            ["Proximity", "(2, 5)", "5", "32", "(5 + 2 × 6/49) / 34 = 0.154262"],
        ],
        [1600, 1400, 1100, 1400, 3860],
    )
    add_equation(
        document,
        "combined_score(17) = (0.158556 + 0.154262) / 2 = 0.156409",
        "The grid formats this value as 15.64%. Rank depends on how this score and the tie-break values compare with the other 48 candidates.",
    )
    add_body(
        document,
        "This example shows why the displayed percentage should be read as a model score based on pooled historical "
        "state outcomes. It is not calibrated here as a promise that number 17 has a 15.64% chance in every future context."
    )

    add_heading(document, "7. Stored objects and import lifecycle", 1)
    add_body(
        document,
        "Prediction information is materialized when a trusted pickle dataset is imported. Draws.load_trusted_pickle() "
        "validates that the loaded object is a Draws instance and calls prepare_predictions(). The method replays the "
        "ordered draw list and attaches one CombinedPrediction object to each Draw through its read-only prediction property."
    )
    add_table(
        document,
        ["Object", "Field", "Meaning"],
        [
            ["NumberPrediction", "number, rank, score", "Candidate identity and final result"],
            ["NumberPrediction", "freshness_score, proximity_score", "The two smoothed component rates"],
            ["NumberPrediction", "gap, left_space, right_space", "Exact raw state used for lookup"],
            ["NumberPrediction", "freshness_support, proximity_support", "Exposure counts E supporting each component"],
            ["CombinedPrediction", "reference_draw_number, target_draw_number", "Prediction timing"],
            ["CombinedPrediction", "numbers, top_numbers", "All 49 ranked records and the leading six"],
            ["CombinedPrediction", "actual_numbers", "Next draw outcome attached for display only"],
            ["Draw", "prediction", "Read-only reference to the precomputed CombinedPrediction"],
        ],
        [2100, 3000, 4260],
    )
    add_body(
        document,
        "The desktop bridge performs a defensive preparation check if predictions are absent, then serializes the recent "
        "prediction history for Electron. Internal component scores and support counts remain available on the Python "
        "objects even though the focused grid currently displays only rank, number, combined score, and outcome marker."
    )

    add_heading(document, "7.1 Complexity", 2)
    add_body(
        document,
        "Let D be the number of stored draws. Learning and scoring inspect exactly 49 candidates per draw, so runtime is "
        "O(49D), conventionally written O(D). Storing a full 49-number prediction after each reference requires O(49D) "
        "records. This eager work is intentional: navigation in the prediction window does not recalculate history."
    )

    add_heading(document, "8. Reading the prediction grid", 1)
    for text in (
        "The first rectangle is rank 1, followed by ranks 2 through 49 across successive rows.",
        "Every rectangle uses the same color; color does not encode score bands or categories.",
        "Each rectangle displays the candidate number, rank, and combined score formatted as a percentage.",
        "For historical references, a neutral check mark identifies a number that appeared in the following recorded draw.",
        "For the latest reference, no outcome marker is available until another draw exists in the dataset.",
        "First, Previous, Next, and Latest navigate precomputed predictions without changing their calculations.",
    ):
        add_list_item(document, text, bullet_id)

    document.add_page_break()
    add_heading(document, "9. Boundaries and reproducibility", 1)
    add_body(
        document,
        "The model is deterministic for a fixed ordered dataset. The prior, exact-state keys, equal component weights, "
        "and tie-break rules fully determine the output. Reimporting the same trusted Draws object in the same order "
        "produces the same scores and ranks."
    )
    add_heading(document, "What the method intentionally does not do", 2)
    for text in (
        "It does not bucket gaps or spacing values into low, medium, or high groups.",
        "It does not infer unobserved future left/right spaces; it carries forward the last observed pair.",
        "It does not train on the target draw before scoring that target.",
        "It does not introduce randomization, hidden weighting, or per-number favoritism.",
        "It does not claim that past lottery structure changes the physical randomness of future draws.",
    ):
        add_list_item(document, text, bullet_id)
    add_callout(
        document,
        "Practical use",
        "Use the grid to compare exact-state empirical rankings and to inspect their historical fit. Treat it as an "
        "analytical view of the supplied draw history, not as certainty or financial advice.",
    )

    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.core_properties.title = "Rand AI Combined Prediction Calculation Guide"
    document.core_properties.subject = "Exact freshness and proximity prediction calculations"
    document.core_properties.author = "Rand AI"
    document.core_properties.keywords = "Rand AI, prediction, freshness, proximity, gap, lottery"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
